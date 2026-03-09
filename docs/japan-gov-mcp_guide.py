#!/usr/bin/env python3
"""
japan-gov-mcp 基本的な使い方ガイド — Word文書生成
"""
import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Inches
from PIL import Image, ImageDraw, ImageFont

SCRIPT_DIR = Path(__file__).resolve().parent
OUT = SCRIPT_DIR / "japan-gov-mcp_usage_guide.docx"


# ── ヘルパー ──

def set_cell_bg(cell, rgb_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tc_pr.append(shd)


def set_font(run, name="游明朝", east_asia=None, size=None, bold=False, color=None):
    run.font.name = name
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), east_asia or name)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for tag, attr in [("begin", "begin"), ("separate", "separate"), ("end", "end")]:
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), attr)
        if tag == "begin":
            run._r.append(el)
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = "PAGE"
            run._r.append(instr)
        run._r.append(el) if tag != "begin" else None


def add_para(doc, text, bold=False, align=None, size=11, font="游明朝"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, name=font, size=size, bold=bold)
    if align:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, name="游ゴシック", east_asia="游ゴシック")
    return h


def add_table(doc, headers, rows, caption=None):
    if caption:
        add_para(doc, caption, bold=True, size=10, font="游ゴシック")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        set_cell_bg(cell, "1B4F72")
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, "游ゴシック", size=9, bold=True, color=(255, 255, 255))

    for ridx, row_data in enumerate(rows):
        row = table.add_row().cells
        for cidx, val in enumerate(row_data):
            row[cidx].text = str(val)
            for p in row[cidx].paragraphs:
                for r in p.runs:
                    set_font(r, "游明朝", size=9)
        if ridx % 2 == 1:
            for c in row:
                set_cell_bg(c, "EBF5FB")

    doc.add_paragraph("")
    return table


def add_code_block(doc, code_text, title=None):
    if title:
        add_para(doc, title, bold=True, size=10, font="游ゴシック")
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        rf.set(qn("w:eastAsia"), "Courier New")
        # gray background
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F3F4")
        rpr.append(shd)


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, "游明朝", size=size)
    return p


# ── メイン ──

def main():
    doc = Document()

    # ページ設定
    for section in doc.sections:
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)

    # スタイル設定
    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(10.5)
    rpr = style._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), "游明朝")

    for hn in ("Heading 1", "Heading 2", "Heading 3"):
        h = doc.styles[hn]
        h.font.name = "游ゴシック"
        hrpr = h._element.get_or_add_rPr()
        hrf = hrpr.find(qn("w:rFonts"))
        if hrf is None:
            hrf = OxmlElement("w:rFonts")
            hrpr.insert(0, hrf)
        hrf.set(qn("w:eastAsia"), "游ゴシック")

    # ヘッダー・フッター
    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("japan-gov-mcp 使い方ガイド")
        set_font(hr, "游ゴシック", size=9, color=(120, 120, 120))
        add_page_number(section.footer.paragraphs[0])

    # ══════════════════════════════════════════
    # 表紙
    # ══════════════════════════════════════════
    add_para(doc, "", size=28)
    add_para(doc, "japan-gov-mcp", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, size=32, font="游ゴシック")
    add_para(doc, "基本的な使い方ガイド", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, size=22, font="游ゴシック")
    add_para(doc, "", size=14)
    add_para(doc, "━" * 35, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "", size=10)
    add_para(doc, "日本の中央省庁APIを統合するMCPサーバー",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_para(doc, "97ツール — 36+ API — APIキー不要で多数利用可能",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_para(doc, "", size=14)
    add_para(doc, "━" * 35, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "", size=10)
    add_para(doc, f"バージョン: v3.5.0  |  更新日: {date.today().isoformat()}",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para(doc, "https://github.com/Rei02061986/japan-gov-mcp",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 目次
    # ══════════════════════════════════════════
    add_heading(doc, "目次", level=1)
    toc = [
        "第1章 概要と特長",
        "第2章 インストールとセットアップ",
        "第3章 APIキーの取得と設定",
        "第4章 ツール一覧（97ツール）",
        "第5章 基本的な使い方",
        "第6章 ユースケース集",
        "第7章 シナリオツール（複合分析）",
        "第8章 トラブルシューティング",
        "付録 APIクレジット・利用規約",
    ]
    for item in toc:
        add_para(doc, f"  {item}", size=11)
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第1章 概要と特長
    # ══════════════════════════════════════════
    add_heading(doc, "第1章 概要と特長", level=1)

    add_para(doc, "japan-gov-mcp は、日本の中央省庁が提供する36以上のAPIを統合した "
             "MCP（Model Context Protocol）サーバーです。Claude Desktop や Claude Code などの "
             "MCPクライアントから、政府統計・法令・気象・企業情報・学術データなどに "
             "自然言語でアクセスできます。")

    add_heading(doc, "1.1 主な特長", level=2)
    features = [
        "97ツール: 統計、企業、気象、法令、地理空間、学術など全分野をカバー",
        "APIキー不要で多数利用可能: 気象、法令、学術検索、BOJ統計、NDB医療データなど",
        "Zodスキーマ検証: 全パラメータに型・範囲制約を適用し、不正入力を即座に拒否",
        "シナリオツール: 複数APIを自動連携して「地域経済分析」「企業調査」等を一発実行",
        "日本語対応: ツール説明・パラメータ説明はすべて日本語",
    ]
    for f in features:
        add_bullet(doc, f)

    add_heading(doc, "1.2 対応クライアント", level=2)
    add_table(doc,
        ["クライアント", "対応状況", "備考"],
        [
            ["Claude Desktop", "対応", "claude_desktop_config.json で設定"],
            ["Claude Code (CLI)", "対応", "設定ファイルまたは --mcp-server で指定"],
            ["MCP Inspector", "対応", "npm run inspect で GUI テスト"],
            ["その他MCP対応アプリ", "対応", "stdio transport をサポートする任意のクライアント"],
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第2章 インストールとセットアップ
    # ══════════════════════════════════════════
    add_heading(doc, "第2章 インストールとセットアップ", level=1)

    add_heading(doc, "2.1 前提条件", level=2)
    add_table(doc,
        ["ソフトウェア", "バージョン", "確認コマンド"],
        [
            ["Node.js", "v18 以上", "node --version"],
            ["npm", "v9 以上", "npm --version"],
            ["Git", "任意", "git --version"],
        ])

    add_heading(doc, "2.2 インストール手順", level=2)

    add_code_block(doc, """
# 1. リポジトリをクローン
git clone https://github.com/Rei02061986/japan-gov-mcp.git
cd japan-gov-mcp

# 2. 依存関係をインストール
npm install

# 3. ビルド
npm run build

# 4. 動作確認（APIキー不要のツールで）
echo '{"jsonrpc":"2.0","id":1,"method":"initialize","params":{"protocolVersion":"2024-11-05","capabilities":{},"clientInfo":{"name":"test","version":"1.0"}}}' | node build/index.js
""", title="ターミナルでの操作")

    add_heading(doc, "2.3 Claude Desktop への登録", level=2)
    add_para(doc, "Claude Desktop の設定ファイル claude_desktop_config.json に以下を追加します。")

    add_code_block(doc, """
{
  "mcpServers": {
    "japan-gov": {
      "command": "node",
      "args": ["/path/to/japan-gov-mcp/build/index.js"],
      "env": {
        "ESTAT_APP_ID": "あなたのe-Stat APIキー"
      }
    }
  }
}
""", title="claude_desktop_config.json")

    add_para(doc, "設定ファイルの場所:", bold=True, size=10)
    add_table(doc,
        ["OS", "パス"],
        [
            ["macOS", "~/Library/Application Support/Claude/claude_desktop_config.json"],
            ["Windows", "%APPDATA%/Claude/claude_desktop_config.json"],
            ["Linux", "~/.config/Claude/claude_desktop_config.json"],
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第3章 APIキーの取得と設定
    # ══════════════════════════════════════════
    add_heading(doc, "第3章 APIキーの取得と設定", level=1)

    add_heading(doc, "3.1 APIキー不要のツール（すぐ使える）", level=2)
    add_para(doc, "以下のカテゴリのツールはAPIキーなしで利用可能です。")
    add_table(doc,
        ["カテゴリ", "主なツール"],
        [
            ["気象・防災", "jma_forecast, jma_earthquake, jma_tsunami, amedas_data, jshis_hazard, flood_depth"],
            ["法令・国会", "law_search, law_keyword_search, kokkai_speeches, kokkai_meetings, pubcomment_list"],
            ["統計ダッシュボード", "dashboard_indicators, dashboard_data"],
            ["BOJ統計", "boj_timeseries, boj_major_statistics"],
            ["NDB医療データ", "ndb_inspection_stats, ndb_items, ndb_areas, ndb_range_labels, ndb_hub_proxy"],
            ["学術検索", "ndl_search, jstage_search, cinii_search, japansearch_search, irdb_search"],
            ["地理空間", "gsi_geocode, gsi_reverse_geocode, geoshape_city, plateau_datasets"],
            ["科学", "soramame_air, geology_at_point, jaxa_collections, agriknowledge_search"],
            ["オープンデータ", "opendata_search, opendata_detail, geospatial_search"],
            ["その他", "safety_overseas, mirasapo_search, traffic_volume, gov_api_catalog"],
        ])

    add_heading(doc, "3.2 APIキーが必要なツール", level=2)
    add_table(doc,
        ["環境変数", "API", "提供元", "取得先URL", "費用"],
        [
            ["ESTAT_APP_ID", "e-Stat（政府統計）", "総務省", "https://www.e-stat.go.jp/api/", "無料"],
            ["HOUJIN_APP_ID", "法人番号", "国税庁", "https://www.houjin-bangou.nta.go.jp/webapi/", "無料"],
            ["GBIZ_TOKEN", "gBizINFO", "経済産業省", "https://info.gbiz.go.jp/hojin/api", "無料"],
            ["EDINET_API_KEY", "EDINET", "金融庁", "https://disclosure2.edinet-fsa.go.jp/", "無料"],
            ["MLIT_DPF_API_KEY", "国交省DPF", "国土交通省", "https://www.mlit-data.jp/", "無料"],
            ["REALESTATE_API_KEY", "不動産情報", "国土交通省", "https://www.reinfolib.mlit.go.jp/", "無料"],
            ["HELLOWORK_API_KEY", "ハローワーク", "厚生労働省", "https://www.hellowork.mhlw.go.jp/", "無料"],
        ])
    add_para(doc, "すべてのAPIキーは無料で取得できます。", bold=True, size=10)

    add_heading(doc, "3.3 環境変数の設定方法", level=2)
    add_code_block(doc, """
# 方法1: .envファイル（推奨）
cp .env.example .env
# .env を編集して各キーを設定

# 方法2: シェルで直接設定
export ESTAT_APP_ID=your_api_key_here

# 方法3: Claude Desktop の env セクションで設定
# (claude_desktop_config.json の "env" に追記)
""", title="設定方法")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第4章 ツール一覧
    # ══════════════════════════════════════════
    add_heading(doc, "第4章 ツール一覧（97ツール）", level=1)

    add_heading(doc, "4.1 統計データ", level=2)
    add_table(doc,
        ["ツール名", "説明", "APIキー"],
        [
            ["estat_search", "e-Stat 統計表検索", "ESTAT_APP_ID"],
            ["estat_meta", "統計表のメタデータ取得", "ESTAT_APP_ID"],
            ["estat_data", "統計データ取得", "ESTAT_APP_ID"],
            ["estat_browse_indicators", "指標一覧ブラウズ", "ESTAT_APP_ID"],
            ["estat_check_availability", "データ利用可否チェック", "ESTAT_APP_ID"],
            ["estat_merger_check", "市区町村合併チェック", "不要"],
            ["estat_compare_municipalities", "自治体比較", "ESTAT_APP_ID"],
            ["estat_time_series", "時系列分析", "ESTAT_APP_ID"],
            ["estat_correlation", "相関分析", "ESTAT_APP_ID"],
            ["estat_session_init", "e-Statセッション初期化", "ESTAT_APP_ID"],
            ["dashboard_indicators", "統計ダッシュボード指標一覧", "不要"],
            ["dashboard_data", "統計ダッシュボードデータ", "不要"],
            ["boj_timeseries", "日銀時系列統計", "不要"],
            ["boj_major_statistics", "日銀主要統計一覧", "不要"],
        ])

    add_heading(doc, "4.2 RESAS（地域経済分析）", level=2)
    add_table(doc,
        ["ツール名", "説明"],
        [
            ["resas_prefectures", "都道府県一覧"],
            ["resas_cities", "市区町村一覧"],
            ["resas_population", "人口推移"],
            ["resas_population_pyramid", "人口ピラミッド"],
            ["resas_industry", "産業構造"],
            ["resas_tourism", "観光データ"],
            ["resas_finance", "財政データ"],
            ["resas_patents", "特許データ"],
        ],
        caption="※ すべて RESAS_API_KEY が必要（無料）")

    add_heading(doc, "4.3 企業情報", level=2)
    add_table(doc,
        ["ツール名", "説明", "APIキー"],
        [
            ["houjin_search", "法人番号検索", "HOUJIN_APP_ID"],
            ["gbiz_search", "gBizINFO企業検索", "GBIZ_TOKEN"],
            ["gbiz_detail", "gBizINFO企業詳細", "GBIZ_TOKEN"],
            ["edinet_documents", "EDINET有価証券報告書", "EDINET_API_KEY"],
        ])

    add_heading(doc, "4.4 気象・防災", level=2)
    add_table(doc,
        ["ツール名", "説明"],
        [
            ["jma_forecast", "天気予報（地域コード指定）"],
            ["jma_overview", "天気概況"],
            ["jma_forecast_week", "週間天気予報"],
            ["jma_typhoon", "台風情報"],
            ["jma_earthquake", "地震情報"],
            ["jma_tsunami", "津波情報"],
            ["amedas_stations", "AMeDAS観測所一覧"],
            ["amedas_data", "AMeDAS観測データ"],
            ["jshis_hazard", "J-SHIS地震ハザード"],
            ["flood_depth", "浸水想定深"],
            ["river_level", "河川水位"],
            ["traffic_volume", "交通量データ（JARTIC）"],
        ],
        caption="※ すべてAPIキー不要")

    add_heading(doc, "4.5 法令・国会", level=2)
    add_table(doc,
        ["ツール名", "説明"],
        [
            ["law_search", "e-Laws法令検索"],
            ["law_data", "法令本文取得"],
            ["law_keyword_search", "法令キーワード検索"],
            ["kokkai_speeches", "国会会議録（発言検索）"],
            ["kokkai_meetings", "国会会議一覧"],
            ["pubcomment_list", "パブリックコメント一覧"],
        ],
        caption="※ すべてAPIキー不要")

    add_heading(doc, "4.6 地理空間・不動産", level=2)
    add_table(doc,
        ["ツール名", "説明", "APIキー"],
        [
            ["gsi_geocode", "住所→座標変換", "不要"],
            ["gsi_reverse_geocode", "座標→住所変換", "不要"],
            ["geoshape_city", "市区町村境界GeoJSON", "不要"],
            ["geoshape_pref", "都道府県境界GeoJSON", "不要"],
            ["plateau_datasets", "PLATEAU 3D都市モデル検索", "不要"],
            ["plateau_citygml", "PLATEAU CityGMLデータ", "不要"],
            ["realestate_transactions", "不動産取引価格", "REALESTATE_API_KEY"],
            ["realestate_landprice", "地価公示・調査", "REALESTATE_API_KEY"],
            ["mlit_dpf_search", "国交省DPFデータ検索", "MLIT_DPF_API_KEY"],
            ["mlit_dpf_catalog", "国交省DPFカタログ", "MLIT_DPF_API_KEY"],
        ])

    doc.add_page_break()

    add_heading(doc, "4.7 学術・文化", level=2)
    add_table(doc,
        ["ツール名", "説明"],
        [
            ["ndl_search", "国立国会図書館検索"],
            ["jstage_search", "J-STAGE論文検索"],
            ["cinii_search", "CiNii学術論文検索"],
            ["japansearch_search", "ジャパンサーチ（文化資源）"],
            ["irdb_search", "IRDB学術機関リポジトリ"],
            ["kkj_search", "建設関連情報検索"],
            ["researchmap_achievements", "researchmap研究者業績"],
            ["agriknowledge_search", "農業知識検索"],
        ],
        caption="※ すべてAPIキー不要")

    add_heading(doc, "4.8 科学・環境", level=2)
    add_table(doc,
        ["ツール名", "説明"],
        [
            ["soramame_air", "大気環境（そらまめくん）"],
            ["geology_legend", "地質図凡例"],
            ["geology_at_point", "地点の地質情報"],
            ["jaxa_collections", "JAXA衛星データカタログ"],
        ],
        caption="※ すべてAPIキー不要")

    add_heading(doc, "4.9 NDB医療データ", level=2)
    add_table(doc,
        ["ツール名", "説明"],
        [
            ["ndb_inspection_stats", "NDB検査統計"],
            ["ndb_items", "NDB検査項目一覧"],
            ["ndb_areas", "NDB地域一覧"],
            ["ndb_range_labels", "NDB基準値ラベル"],
            ["ndb_hub_proxy", "NDB Hubプロキシ"],
        ],
        caption="※ すべてAPIキー不要")

    add_heading(doc, "4.10 その他", level=2)
    add_table(doc,
        ["ツール名", "説明", "APIキー"],
        [
            ["safety_overseas", "海外安全情報", "不要"],
            ["hellowork_search", "ハローワーク求人検索", "HELLOWORK_API_KEY"],
            ["mirasapo_search", "ミラサポplus補助金検索", "不要"],
            ["mirasapo_detail", "ミラサポplus補助金詳細", "不要"],
            ["mirasapo_categories", "ミラサポplusカテゴリ", "不要"],
            ["mirasapo_regions", "ミラサポplus地域", "不要"],
            ["opendata_search", "data.go.jpオープンデータ検索", "不要"],
            ["opendata_detail", "data.go.jpデータ詳細", "不要"],
            ["geospatial_search", "G空間データ検索", "不要"],
            ["geospatial_dataset", "G空間データセット詳細", "不要"],
            ["geospatial_organizations", "G空間データ提供機関", "不要"],
            ["msil_layers", "MSIL地図レイヤー一覧", "不要"],
            ["msil_features", "MSILフィーチャー取得", "不要"],
            ["odpt_railway_timetable", "鉄道時刻表", "不要"],
            ["odpt_bus_timetable", "バス時刻表", "不要"],
            ["gov_api_catalog", "政府API一覧", "不要"],
            ["gov_cross_search", "横断検索", "不要"],
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第5章 基本的な使い方
    # ══════════════════════════════════════════
    add_heading(doc, "第5章 基本的な使い方", level=1)

    add_para(doc, "MCPサーバーは Claude Desktop や Claude Code から自動的に呼び出されます。"
             "ユーザーは自然言語で質問するだけで、Claudeが適切なツールを選択してデータを取得します。")

    add_heading(doc, "5.1 自然言語での利用例", level=2)
    add_para(doc, "Claude に以下のように聞くだけで、適切なツールが自動的に呼ばれます。")

    examples = [
        ("「東京の明日の天気は？」", "→ jma_forecast が呼ばれ、東京地方の天気予報を返します"),
        ("「少子化に関する統計を探して」", "→ estat_search で出生率・婚姻率等の統計表を検索します"),
        ("「トヨタ自動車の法人番号を調べて」", "→ houjin_search で法人番号を検索します"),
        ("「著作権法の条文を見せて」", "→ law_keyword_search + law_data で法令を取得します"),
        ("「渋谷区の地震ハザード情報は？」", "→ gsi_geocode + jshis_hazard で地震動予測を返します"),
        ("「BMIの全国検査統計を見せて」", "→ ndb_inspection_stats でNDB医療データを取得します"),
        ("「AIに関する最新の論文を検索」", "→ jstage_search + cinii_search で学術論文を横断検索します"),
    ]
    for question, answer in examples:
        p = add_para(doc, "", size=10.5)
        run_q = p.add_run(question)
        set_font(run_q, "游ゴシック", size=10.5, bold=True)
        p.add_run("\n")
        run_a = p.add_run(answer)
        set_font(run_a, "游明朝", size=10)

    add_heading(doc, "5.2 直接ツール指定（上級者向け）", level=2)
    add_para(doc, "Claude Code などで直接ツール名を指定して呼び出すこともできます。")

    add_code_block(doc, """
# 天気予報を取得
tool: jma_forecast
args: { "areaCode": "130000" }

# e-Stat でテーブル検索
tool: estat_search
args: { "keyword": "人口動態 出生", "limit": 10 }

# 法令キーワード検索
tool: law_keyword_search
args: { "keyword": "個人情報保護", "limit": 5 }
""", title="ツール直接呼び出しの例")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第6章 ユースケース集
    # ══════════════════════════════════════════
    add_heading(doc, "第6章 ユースケース集", level=1)

    add_heading(doc, "6.1 地域分析", level=2)
    add_para(doc, "「長門市の産業構造を分析して、他の温泉観光地と比較して」", bold=True, size=11)
    add_para(doc, "Claude は以下のツールを組み合わせて自動的に分析します:")
    add_bullet(doc, "estat_search → 経済センサスのテーブルIDを特定")
    add_bullet(doc, "estat_data → 長門市・比較自治体のデータを取得")
    add_bullet(doc, "dashboard_data → 統計ダッシュボードで補完")
    add_bullet(doc, "結果を表やグラフにまとめて報告")

    add_heading(doc, "6.2 企業調査", level=2)
    add_para(doc, "「ソフトバンクグループの子会社一覧と直近の有価証券報告書を調べて」", bold=True, size=11)
    add_bullet(doc, "houjin_search → 法人番号を検索")
    add_bullet(doc, "gbiz_detail → gBizINFOで企業詳細を取得")
    add_bullet(doc, "edinet_documents → EDINET有価証券報告書を検索")

    add_heading(doc, "6.3 防災リスク評価", level=2)
    add_para(doc, "「東京都港区の地震・浸水リスクを評価して」", bold=True, size=11)
    add_bullet(doc, "gsi_geocode → 住所から座標を取得")
    add_bullet(doc, "jshis_hazard → 地震ハザード情報を取得")
    add_bullet(doc, "flood_depth → 浸水想定深を取得")
    add_bullet(doc, "scenario_disaster_risk_assessment → 総合的なリスク評価")

    add_heading(doc, "6.4 学術調査", level=2)
    add_para(doc, "「機械学習を用いた医療診断に関する日本の研究動向を調べて」", bold=True, size=11)
    add_bullet(doc, "jstage_search → J-STAGEで論文検索")
    add_bullet(doc, "cinii_search → CiNiiで論文検索")
    add_bullet(doc, "scenario_academic_trend → 研究トレンド分析")

    add_heading(doc, "6.5 日銀・経済データ", level=2)
    add_para(doc, "「最新のドル円為替レートと消費者物価指数の推移を見せて」", bold=True, size=11)
    add_bullet(doc, "boj_timeseries → 日銀時系列データ（為替・金利）")
    add_bullet(doc, "dashboard_data → 統計ダッシュボード（CPI等）")
    add_bullet(doc, "scenario_national_economy_summary → 経済サマリー")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第7章 シナリオツール
    # ══════════════════════════════════════════
    add_heading(doc, "第7章 シナリオツール（複合分析）", level=1)

    add_para(doc, "シナリオツールは、複数のAPIを自動的に組み合わせて総合的な分析結果を返します。"
             "1回の呼び出しで複数のデータソースを横断的に分析できます。")

    add_table(doc,
        ["ツール名", "説明", "使用API"],
        [
            ["scenario_regional_health_economy", "地域の健康・経済分析", "e-Stat, NDB, RESAS"],
            ["scenario_labor_demand_supply", "労働需給分析", "RESAS, e-Stat, HelloWork"],
            ["scenario_corporate_intelligence", "企業インテリジェンス", "法人番号, gBiz, EDINET"],
            ["scenario_disaster_risk_assessment", "災害リスク総合評価", "JMA, J-SHIS, 浸水ナビ"],
            ["scenario_academic_trend", "学術トレンド分析", "J-STAGE, CiNii, IRDB"],
            ["scenario_academic_trend_by_topics", "トピック別学術トレンド", "J-STAGE, CiNii"],
            ["scenario_realestate_demographics", "不動産×人口分析", "不動産, e-Stat, RESAS"],
            ["scenario_regional_economy_full", "地域経済フル分析", "e-Stat, RESAS, Dashboard"],
            ["scenario_national_economy_summary", "全国経済サマリー", "BOJ, Dashboard"],
        ],
        caption="シナリオツール一覧")

    add_para(doc, "使用例:", bold=True, size=10)
    add_code_block(doc, """
# 地域の健康・経済分析
tool: scenario_regional_health_economy
args: { "prefectureCode": "13" }
→ 東京都の健康指標・経済指標を複合分析

# 災害リスク評価
tool: scenario_disaster_risk_assessment
args: { "lat": 35.6762, "lon": 139.6503, "address": "東京都渋谷区" }
→ 地震ハザード + 浸水リスク + 気象情報を統合評価
""")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 第8章 トラブルシューティング
    # ══════════════════════════════════════════
    add_heading(doc, "第8章 トラブルシューティング", level=1)

    add_table(doc,
        ["症状", "原因", "対策"],
        [
            ["「APIキーが設定されていません」", "環境変数が未設定", ".envファイルまたはclaude_desktop_config.jsonのenvセクションにキーを設定"],
            ["データが取得できない", "外部APIのメンテナンス", "しばらく待ってから再試行。APIの稼働状況を確認"],
            ["タイムアウト", "外部APIの応答遅延", "一部のAPI（そらまめくん等）は応答に10-20秒かかることがある"],
            ["「Zod validation error」", "パラメータの型・範囲エラー", "ツールのパラメータ説明を確認（数値範囲、文字列フォーマット等）"],
            ["Claude Desktopで認識されない", "設定ファイルの書式エラー", "JSONの書式を確認。パスにスペースがある場合はエスケープ"],
            ["ビルドエラー", "Node.jsバージョンが古い", "Node.js v18以上を使用"],
        ],
        caption="よくある問題と対策")

    add_heading(doc, "8.1 ログの確認", level=2)
    add_para(doc, "japan-gov-mcp は全ツール呼び出しをJSONL形式でログに記録します。")
    add_code_block(doc, """
# ログファイルの場所
~/.japan-gov-mcp/logs/

# 最も呼ばれたツール
cat ~/.japan-gov-mcp/logs/*.jsonl | jq -r .tool | sort | uniq -c | sort -rn

# エラーのみ抽出
cat ~/.japan-gov-mcp/logs/*.jsonl | jq 'select(.status=="error")'

# ログ無効化
export JAPAN_GOV_MCP_LOG=false
""")

    add_heading(doc, "8.2 MCP Inspector でのテスト", level=2)
    add_para(doc, "MCP Inspector を使うと、ブラウザ上のGUIで各ツールを個別にテストできます。")
    add_code_block(doc, """
# MCP Inspector を起動
npm run inspect

# ブラウザで http://localhost:5173 を開く
# ツール一覧からテストしたいツールを選択し、パラメータを入力して実行
""")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 付録
    # ══════════════════════════════════════════
    add_heading(doc, "付録 APIクレジット・利用規約", level=1)

    add_para(doc, "本MCPサーバーが使用する各APIのデータは、それぞれの提供元の利用規約に従います。"
             "以下は主要なクレジット表記です。")

    credits = [
        ["e-Stat", "政府統計総合窓口(e-Stat)のAPI機能を使用。サービスの内容は国によって保証されたものではありません。"],
        ["統計ダッシュボード", "統計ダッシュボードのAPI機能を使用。サービスの内容は国によって保証されたものではありません。"],
        ["BOJ時系列統計", "日本銀行時系列統計データ検索サイトのAPI機能を使用。サービスの内容は日本銀行によって保証されたものではありません。"],
        ["法人番号", "国税庁法人番号システムWeb-API機能を利用。サービスの内容は国税庁によって保証されたものではありません。"],
        ["不動産情報ライブラリ", "国土交通省不動産情報ライブラリのAPI機能を使用。提供情報の正確性、完全性等は保証されていません。"],
    ]
    add_table(doc, ["API", "クレジット表記"], credits, caption="必要なクレジット表記")

    add_para(doc, "", size=10)
    add_para(doc, "免責事項", bold=True, size=12, font="游ゴシック")
    add_para(doc,
        "本プロジェクトは日本政府・各省庁とは無関係の非公式ツールです。"
        "取得データの正確性・完全性・可用性は保証しません。"
        "政府APIは予告なく変更・停止される場合があります。"
        "重要な意思決定には必ず公式情報源を確認してください。", size=10)

    add_para(doc, "", size=10)
    add_para(doc, "━" * 35, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"japan-gov-mcp v3.5.0 使い方ガイド — {date.today().isoformat()}",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    add_para(doc, "https://github.com/Rei02061986/japan-gov-mcp",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    # 保存
    doc.save(str(OUT))
    print(f"Generated: {OUT}")
    print(f"File size: {OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
